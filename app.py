import os
import re
import logging
from copy import deepcopy
from pathlib import Path
from typing import Optional

from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, Depends, Security
from fastapi.security.api_key import APIKeyHeader
from pydantic import BaseModel, Field
from openai import OpenAI

load_dotenv()

# -------------------------
# LOGGING
# -------------------------

logging.basicConfig(
    level=os.getenv("LOG_LEVEL", "INFO"),
    format="%(asctime)s | %(levelname)s | %(name)s | %(message)s"
)
logger = logging.getLogger("isdm_rag_api")

# -------------------------
# APP
# -------------------------

app = FastAPI(
    title="ISDM RAG API",
    version="0.4.0",
    servers=[{"url": "https://isdm-rag.onrender.com"}]
)

# -------------------------
# SEGURIDAD API KEY
# -------------------------

API_KEY = os.getenv("ISDM_API_KEY")
API_KEY_NAME = "X-API-KEY"

api_key_header = APIKeyHeader(
    name=API_KEY_NAME,
    auto_error=False
)


def require_api_key(api_key: str = Security(api_key_header)):
    if not API_KEY:
        raise HTTPException(status_code=500, detail="ISDM_API_KEY not configured")

    if not api_key or api_key != API_KEY:
        raise HTTPException(status_code=401, detail="Invalid or missing API key")

    return api_key


# -------------------------
# CONFIG / CLIENT
# -------------------------

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    raise RuntimeError("Missing env var: OPENAI_API_KEY")

client = OpenAI(api_key=OPENAI_API_KEY)

VECTOR_STORE_NAME = os.environ.get("VECTOR_STORE_NAME", "ISDM_2025_C1")

OUTPUT_ROOT_DIR = Path(
    os.environ.get(
        "OUTPUT_ROOT_DIR",
        "./generated_docs"
    )
)

MODEL_NAME = os.environ.get("OPENAI_MODEL", "gpt-4.1-mini")


def get_vector_store_id() -> str:
    stores = client.vector_stores.list(limit=100)
    for s in stores.data:
        if getattr(s, "name", None) == VECTOR_STORE_NAME:
            return s.id
    created = client.vector_stores.create(name=VECTOR_STORE_NAME)
    return created.id


VECTOR_STORE_ID = get_vector_store_id()

# -------------------------
# BASIC ENDPOINTS (PÚBLICOS)
# -------------------------


@app.get("/")
def root():
    return {
        "ok": True,
        "service": "ISDM_RAG",
        "version": app.version,
        "vector_store_name": VECTOR_STORE_NAME,
        "vector_store_id": VECTOR_STORE_ID,
        "output_root_dir": str(OUTPUT_ROOT_DIR),
    }


@app.get("/health")
def health():
    return {"ok": True}


# -------------------------
# MODELOS
# -------------------------


class SearchRequest(BaseModel):
    query: str = Field(..., min_length=2)
    k: int = Field(8, ge=1, le=20)
    anio: Optional[str] = None
    materia: Optional[str] = None
    unidad: Optional[str] = None
    unidad_num: Optional[str] = None
    tipo: Optional[str] = None


class SearchResult(BaseModel):
    text: str
    score: float
    file_id: str
    attributes: Optional[dict] = None


class SearchResponse(BaseModel):
    vector_store_id: str
    results: list[SearchResult]


class ImproveRequest(BaseModel):
    objetivo: str = Field(..., min_length=5)
    query: str = Field(..., min_length=2)
    anio: Optional[str] = None
    materia: Optional[str] = None
    unidad: Optional[str] = None
    k: int = Field(10, ge=3, le=20)
    unidad_num: Optional[str] = None
    tipo: Optional[str] = None

    modo_fuentes: str = Field(
        "optional",
        pattern="^(required|optional|none)$"
    )

    nivel: str = "Superior ISDM"
    formato_salida: str = "plan_clase"
    tono: str = "profesional_cercano"
    restricciones: Optional[str] = None

    # nuevos campos opcionales, no rompen compatibilidad
    cantidad_documentos: int = Field(1, ge=1, le=20)
    output_subdir: Optional[str] = None


class ImproveResponse(BaseModel):
    used_sources: list[dict]
    output: str


# -------------------------
# HELPERS
# -------------------------


def sanitize_filename(value: Optional[str], fallback: str) -> str:
    text = (value or "").strip()
    if not text:
        text = fallback

    replacements = {
        "á": "a", "à": "a", "ä": "a", "â": "a",
        "é": "e", "è": "e", "ë": "e", "ê": "e",
        "í": "i", "ì": "i", "ï": "i", "î": "i",
        "ó": "o", "ò": "o", "ö": "o", "ô": "o",
        "ú": "u", "ù": "u", "ü": "u", "û": "u",
        "ñ": "n",
        "Á": "A", "À": "A", "Ä": "A", "Â": "A",
        "É": "E", "È": "E", "Ë": "E", "Ê": "E",
        "Í": "I", "Ì": "I", "Ï": "I", "Î": "I",
        "Ó": "O", "Ò": "O", "Ö": "O", "Ô": "O",
        "Ú": "U", "Ù": "U", "Ü": "U", "Û": "U",
        "Ñ": "N",
    }

    for src, dst in replacements.items():
        text = text.replace(src, dst)

    text = re.sub(r"[^\w\s-]", "", text, flags=re.UNICODE)
    text = re.sub(r"[-\s]+", "_", text.strip())
    return text or fallback


def normalize_materia_folder(materia: Optional[str]) -> str:
    raw = (materia or "").strip().lower()

    aliases = {
        "taller de programación": "TallerDeProgramacion",
        "taller de programacion": "TallerDeProgramacion",
        "taller_programacion": "TallerDeProgramacion",
        "tallerdeprogramacion": "TallerDeProgramacion",

        "práctica profesional ii": "PracticaProfesionalII",
        "practica profesional ii": "PracticaProfesionalII",
        "práctica profesional 2": "PracticaProfesionalII",
        "practica profesional 2": "PracticaProfesionalII",
        "pp2": "PracticaProfesionalII",
        "practicaprofesionalii": "PracticaProfesionalII",
    }

    if raw in aliases:
        return aliases[raw]

    cleaned = sanitize_filename(materia, "General")
    parts = [p for p in cleaned.split("_") if p]
    return "".join(p[:1].upper() + p[1:] for p in parts) or "General"


def detect_file_prefix(req: ImproveRequest) -> str:
    text = f"{req.query} {req.objetivo} {req.formato_salida} {req.tipo or ''}".lower()

    patterns = [
        (r"\bprimera clase\b|\bclase 1\b|\bclase uno\b", "Clase_1"),
        (r"\bsegunda clase\b|\bclase 2\b|\bclase dos\b", "Clase_2"),
        (r"\btercera clase\b|\bclase 3\b|\bclase tres\b", "Clase_3"),
        (r"\bcuarta clase\b|\bclase 4\b|\bclase cuatro\b", "Clase_4"),
        (r"\bquinta clase\b|\bclase 5\b|\bclase cinco\b", "Clase_5"),
        (r"\bsexta clase\b|\bclase 6\b|\bclase seis\b", "Clase_6"),
        (r"\bséptima clase\b|\bseptima clase\b|\bclase 7\b", "Clase_7"),
        (r"\boctava clase\b|\bclase 8\b", "Clase_8"),
        (r"\bnovena clase\b|\bclase 9\b", "Clase_9"),
        (r"\bdécima clase\b|\bdecima clase\b|\bclase 10\b", "Clase_10"),
    ]

    for pattern, prefix in patterns:
        if re.search(pattern, text):
            return prefix

    if "diagnostico" in text or "diagnóstico" in text:
        return "Diagnostico"

    if "cronograma" in text:
        return "Cronograma"

    if "planificación" in text or "planificacion" in text:
        return "Planificacion"

    if "clase" in text:
        return "Clase"

    return "Documento"


def detect_subfolder(req: ImproveRequest) -> Optional[str]:
    if req.output_subdir:
        return sanitize_filename(req.output_subdir, "subdir")

    text = f"{req.query} {req.objetivo} {req.unidad or ''} {req.tipo or ''}".lower()

    if "diagnostico" in text or "diagnóstico" in text:
        if req.materia and "taller" in req.materia.lower():
            return "diagnostico_programacion"
        return "diagnostico"

    return None


def extract_explicit_class_numbers(text: str) -> list[int]:
    found = set()

    numeric = re.findall(r"\bclases?\s*(\d+)\s*(?:,|y|-|a)?\s*(\d+)?\b", text.lower())
    for a, b in numeric:
        if a:
            found.add(int(a))
        if b:
            found.add(int(b))

    standalone = re.findall(r"\bclase\s*(\d+)\b", text.lower())
    for n in standalone:
        found.add(int(n))

    word_map = {
        "primera": 1, "segunda": 2, "tercera": 3, "cuarta": 4, "quinta": 5,
        "sexta": 6, "septima": 7, "séptima": 7, "octava": 8, "novena": 9, "decima": 10, "décima": 10,
    }

    for word, num in word_map.items():
        if re.search(rf"\b{word}\b", text.lower()):
            found.add(num)

    return sorted(n for n in found if 1 <= n <= 50)


def infer_multi_class_requests(req: ImproveRequest) -> list[ImproveRequest]:
    """
    Mantiene compatibilidad:
    - si no detecta multi-documento => devuelve [req]
    - si detecta varias clases => devuelve N requests derivados
    """
    if req.cantidad_documentos > 1:
        generated = []
        class_nums = extract_explicit_class_numbers(f"{req.query} {req.objetivo}")

        if not class_nums:
            class_nums = list(range(1, req.cantidad_documentos + 1))

        for n in class_nums[:req.cantidad_documentos]:
            cloned = req.model_copy(deep=True)
            cloned.query = force_query_for_class(req.query, n)
            cloned.objetivo = force_goal_for_class(req.objetivo, n)
            cloned.cantidad_documentos = 1
            generated.append(cloned)

        return generated

    text = f"{req.query} {req.objetivo}".lower()

    if any(term in text for term in ["dos clases", "2 clases", "ambas clases", "las dos clases"]):
        class_nums = extract_explicit_class_numbers(text)
        if not class_nums:
            class_nums = [1, 2]

        generated = []
        for n in class_nums:
            cloned = req.model_copy(deep=True)
            cloned.query = force_query_for_class(req.query, n)
            cloned.objetivo = force_goal_for_class(req.objetivo, n)
            generated.append(cloned)
        return generated

    class_nums = extract_explicit_class_numbers(text)
    if len(class_nums) > 1:
        generated = []
        for n in class_nums:
            cloned = req.model_copy(deep=True)
            cloned.query = force_query_for_class(req.query, n)
            cloned.objetivo = force_goal_for_class(req.objetivo, n)
            generated.append(cloned)
        return generated

    return [req]


def force_query_for_class(original: str, class_number: int) -> str:
    return f"{original}\n\nEste documento debe corresponder específicamente a la Clase {class_number}."


def force_goal_for_class(original: str, class_number: int) -> str:
    return f"{original}\n\nGenerar específicamente el documento correspondiente a la Clase {class_number}."


def build_generation_prompt(req: ImproveRequest, context_body: str, has_context: bool) -> str:
    fuentes_texto = (
        "Usar únicamente el contexto ISDM recuperado para afirmaciones institucionales. "
        "Si falta evidencia puntual, indicarlo explícitamente."
        if has_context
        else "No hay contexto ISDM disponible. Generar una propuesta general y aclararlo explícitamente."
    )

    return f"""
Objetivo del encargo:
{req.objetivo}

Pedido del usuario:
{req.query}

Parámetros de generación:
- Nivel: {req.nivel}
- Formato de salida: {req.formato_salida}
- Tono: {req.tono}
- Materia: {req.materia or "No especificada"}
- Año: {req.anio or "No especificado"}
- Unidad: {req.unidad or "No especificada"}
- Unidad numérica: {req.unidad_num or "No especificada"}
- Tipo: {req.tipo or "No especificado"}
- Restricciones: {req.restricciones or "Ninguna"}

Reglas de respuesta:
- Responder en español.
- Entregar una salida directamente utilizable.
- Mantener estilo profesional, claro y operativo.
- No inventar respaldo institucional.
- {fuentes_texto}
- Si se usaron fuentes ISDM, cerrar con una sección breve llamada "Fuentes ISDM recuperadas:".
- Si no se encontraron o no se usaron fuentes, aclararlo explícitamente al final.
- Si el pedido corresponde a una clase, estructurar la salida como planificación docente profesional.
- Si el pedido menciona duración u horas, respetarlas en la planificación.
- No incluir explicaciones meta ni texto conversacional.
- Escribir el contenido final listo para pasar a Word.

Contexto recuperado:
{context_body}
""".strip()


def parse_text_to_docx(document, output_text: str):
    blocks = [b.strip() for b in output_text.split("\n\n") if b.strip()]

    for block in blocks:
        lines = [ln.strip() for ln in block.splitlines() if ln.strip()]
        if not lines:
            continue

        first = lines[0]

        if first.startswith("#"):
            level = min(first.count("#"), 3)
            heading_text = first.lstrip("#").strip() or "Sección"
            document.add_heading(heading_text, level=level)
            for ln in lines[1:]:
                if re.match(r"^[-•*]\s+", ln):
                    document.add_paragraph(re.sub(r"^[-•*]\s+", "", ln), style="List Bullet")
                else:
                    document.add_paragraph(ln)
            continue

        if re.match(r"^\d+[\.\)]\s+", first):
            document.add_heading(first, level=2)
            for ln in lines[1:]:
                if re.match(r"^[-•*]\s+", ln):
                    document.add_paragraph(re.sub(r"^[-•*]\s+", "", ln), style="List Bullet")
                else:
                    document.add_paragraph(ln)
            continue

        if len(lines) == 1 and len(first) <= 110 and not first.endswith("."):
            document.add_heading(first, level=2)
            continue

        for ln in lines:
            if re.match(r"^[-•*]\s+", ln):
                document.add_paragraph(re.sub(r"^[-•*]\s+", "", ln), style="List Bullet")
            else:
                document.add_paragraph(ln)


def save_output_document(output_text: str, req: ImproveRequest) -> Path:
    """
    Mantiene la lógica usada en Taller:
    - carpeta base: OUTPUT_ROOT_DIR
    - subcarpeta por materia normalizada
    - subcarpeta opcional detectada por contexto
    - nombre base a partir del pedido
    - si existe, agrega sufijo _2, _3, etc.
    """
    materia_dir_name = normalize_materia_folder(req.materia)
    target_dir = OUTPUT_ROOT_DIR / materia_dir_name

    subfolder = detect_subfolder(req)
    if subfolder:
        target_dir = target_dir / subfolder

    target_dir.mkdir(parents=True, exist_ok=True)

    prefix = detect_file_prefix(req)

    candidate_docx = target_dir / f"{prefix}.docx"
    candidate_md = target_dir / f"{prefix}.md"

    if candidate_docx.exists() or candidate_md.exists():
        suffix = 2
        while True:
            alt_docx = target_dir / f"{prefix}_{suffix}.docx"
            alt_md = target_dir / f"{prefix}_{suffix}.md"
            if not alt_docx.exists() and not alt_md.exists():
                candidate_docx = alt_docx
                candidate_md = alt_md
                break
            suffix += 1

    try:
        from docx import Document  # type: ignore

        document = Document()
        parse_text_to_docx(document, output_text)
        document.save(candidate_docx)
        logger.info("Archivo DOCX guardado en: %s", candidate_docx)
        return candidate_docx

    except Exception as e:
        logger.exception("Fallo guardado DOCX. Se intentará MD. Error: %s", e)
        candidate_md.write_text(output_text, encoding="utf-8")
        logger.info("Archivo MD guardado en: %s", candidate_md)
        return candidate_md


def extract_text_from_search_item(item) -> str:
    txt = ""
    content = getattr(item, "content", None)

    if isinstance(content, list):
        parts = []
        for c in content:
            if isinstance(c, dict):
                if "text" in c and c["text"]:
                    parts.append(str(c["text"]))
                elif "content" in c and c["content"]:
                    parts.append(str(c["content"]))
            else:
                c_text = getattr(c, "text", None)
                if c_text:
                    parts.append(str(c_text))
                else:
                    c_content = getattr(c, "content", None)
                    if c_content:
                        parts.append(str(c_content))
        txt = "\n".join(parts).strip()

    elif isinstance(content, dict):
        txt = str(content.get("text") or content.get("content") or "").strip()

    elif content is not None:
        txt = str(content).strip()

    return txt


def build_filters(req: SearchRequest):
    filter_clauses = []

    if req.anio:
        filter_clauses.append({
            "type": "eq",
            "key": "anio",
            "value": req.anio
        })

    if req.materia:
        filter_clauses.append({
            "type": "eq",
            "key": "materia",
            "value": req.materia
        })

    if req.unidad:
        filter_clauses.append({
            "type": "eq",
            "key": "unidad",
            "value": req.unidad
        })

    if req.unidad_num:
        filter_clauses.append({
            "type": "eq",
            "key": "unidad_num",
            "value": req.unidad_num
        })

    if req.tipo:
        filter_clauses.append({
            "type": "eq",
            "key": "tipo",
            "value": req.tipo
        })

    if len(filter_clauses) == 1:
        return filter_clauses[0]

    if len(filter_clauses) > 1:
        return {
            "type": "and",
            "filters": filter_clauses
        }

    return None


def build_context_and_sources(req: ImproveRequest) -> tuple[list[dict], str]:
    sources: list[dict] = []
    context = ""

    if req.modo_fuentes == "none":
        return sources, context

    sreq = SearchRequest(
        query=req.query,
        k=req.k,
        anio=req.anio,
        materia=req.materia,
        unidad=req.unidad,
        unidad_num=req.unidad_num,
        tipo=req.tipo
    )

    sresp: SearchResponse = search(sreq, API_KEY)

    context_blocks = []
    for i, r in enumerate(sresp.results, start=1):
        meta = r.attributes or {}
        sources.append({
            "rank": i,
            "score": r.score,
            "source_name": meta.get("source_name"),
            "source_path": meta.get("source_path"),
            "anio": meta.get("anio"),
            "materia": meta.get("materia"),
            "unidad": meta.get("unidad"),
            "unidad_num": meta.get("unidad_num"),
            "tipo": meta.get("tipo"),
            "file_id": r.file_id,
        })
        context_blocks.append(
            f"[FUENTE {i}] "
            f"{meta.get('source_name', '(sin nombre)')} | "
            f"{meta.get('materia', '?')} {meta.get('anio', '?')} {meta.get('unidad', '')}\n"
            f"{r.text}\n"
        )

    context = "\n\n".join(context_blocks)

    if req.modo_fuentes == "required" and not sresp.results:
        return [], ""

    return sources, context


def generate_single_output(req: ImproveRequest, context: str) -> str:
    has_context = bool(context.strip())

    if has_context:
        rule = (
            "REGLA: Usá el contexto recuperado como base para cualquier afirmación institucional. "
            "No inventes normativa, documentos ni atribuciones al ISDM."
        )
        context_body = context
    else:
        rule = (
            "REGLA: No hay contexto ISDM recuperado. Generá una propuesta general, "
            "aclarando explícitamente que fue elaborada sin respaldo ISDM encontrado."
        )
        context_body = "[SIN CONTEXTO ISDM]"

    prompt = build_generation_prompt(
        req=req,
        context_body=context_body,
        has_context=has_context
    )

    resp = client.responses.create(
        model=MODEL_NAME,
        input=[
            {"role": "system", "content": rule},
            {"role": "user", "content": prompt},
        ],
    )

    return getattr(resp, "output_text", None) or str(resp)


# -------------------------
# SEARCH (PROTEGIDO)
# -------------------------


@app.post("/search", response_model=SearchResponse)
def search(req: SearchRequest, _: str = Depends(require_api_key)):
    filters = build_filters(req)

    try:
        resp = client.vector_stores.search(
            vector_store_id=VECTOR_STORE_ID,
            query=req.query,
            max_num_results=req.k,
            filters=filters
        )
    except Exception as e:
        logger.exception("Vector store search failed")
        raise HTTPException(status_code=500, detail=f"Vector store search failed: {e}")

    out = []
    for item in resp.data:
        txt = extract_text_from_search_item(item)

        out.append(SearchResult(
            text=txt,
            score=float(getattr(item, "score", 0.0)),
            file_id=str(getattr(item, "file_id", "")),
            attributes=getattr(item, "attributes", None)
        ))

    return SearchResponse(vector_store_id=VECTOR_STORE_ID, results=out)


# -------------------------
# IMPROVE FOR 2026 (PROTEGIDO)
# -------------------------


@app.post("/improve_2026", response_model=ImproveResponse)
def improve_2026(req: ImproveRequest, _: str = Depends(require_api_key)):
    logger.info(
        "Nueva solicitud improve_2026 | materia=%s | unidad=%s | tipo=%s",
        req.materia, req.unidad, req.tipo
    )

    expanded_requests = infer_multi_class_requests(req)
    logger.info("Documentos a generar: %s", len(expanded_requests))

    # Contexto y fuentes se recuperan una vez con el request original
    sources, context = build_context_and_sources(req)

    if req.modo_fuentes == "required" and not context.strip():
        return ImproveResponse(
            used_sources=[],
            output="No encontré fuentes ISDM con esos filtros."
        )

    outputs: list[str] = []
    saved_paths: list[str] = []

    for index, subreq in enumerate(expanded_requests, start=1):
        try:
            logger.info("Generando documento %s/%s", index, len(expanded_requests))
            output_text = generate_single_output(subreq, context)

            saved_path = save_output_document(output_text=output_text, req=subreq)
            saved_paths.append(str(saved_path))

            outputs.append(
                f"=== DOCUMENTO {index} ===\n\n{output_text}\n\n[Archivo guardado en: {saved_path}]"
            )

        except Exception as e:
            logger.exception("Error generando documento %s", index)
            outputs.append(
                f"=== DOCUMENTO {index} ===\n\nNo se pudo generar este documento.\nDetalle técnico: {e}"
            )

    final_output = "\n\n".join(outputs)

    if saved_paths:
        final_output += "\n\nArchivos guardados:\n" + "\n".join(f"- {p}" for p in saved_paths)

    return ImproveResponse(
        used_sources=sources,
        output=final_output
    )